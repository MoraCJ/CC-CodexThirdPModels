#!/usr/bin/env python3
"""Build the unified Claude + Codex proxy runbook in Markdown and DOCX."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
MD_PATH = OUT_DIR / "claude-codex-unified-proxy-runbook.md"
DOCX_PATH = OUT_DIR / "claude-codex-unified-proxy-runbook.docx"

TITLE = "Claude + Codex 统一本机代理 Handoff Runbook"
SUBTITLE = "一个本机 HTTPS 代理同时服务 Claude Code Desktop/CLI 与 Codex CLI/App"
DATE = "2026-05-14"

REMOTE_HOME = "/Users/corptest"
REMOTE_ROOT = f"{REMOTE_HOME}/Documents/Codex/claude-code-app-api"
PROXY_DIR = f"{REMOTE_ROOT}/claude-local-proxy"


SUMMARY = [
    "已把 Codex 代理逻辑合进 Claude 本机 HTTPS 代理，远端 Mac 当前只保留 `https://127.0.0.1:38443` 作为统一入口。",
    "Claude 仍使用 Claude 槽位名，由代理映射到真实上游模型；Codex 直接使用真实模型名，通过 profile 切换。",
    "旧 Codex 独立代理端口 `38444` 已停止监听，旧文件保留作为回滚参考。",
    "验证结果显示 Claude `/v1/messages`、Codex `/v1/responses`、Codex profile 与 tool call 均可用。",
    "本文不包含真实 API key、密码或私钥；需要配置时统一使用 `<ARK_API_KEY>` 占位。",
]


ARCHITECTURE = """Claude Desktop / Claude Code CLI
  POST /v1/messages, /v1/messages/count_tokens
                 \\
                  \\  https://127.0.0.1:38443
                   \\ LaunchAgent + local HTTPS cert
                   / claude-local-proxy/server.js
                  /
Codex CLI / App
  POST /v1/responses

Claude path: passthrough to https://ark.cn-beijing.volces.com/api/coding
Codex path: Responses API -> Chat Completions -> https://ark.cn-beijing.volces.com/api/coding/v3/chat/completions -> Responses API
"""


INVENTORY = [
    ["目标 Mac", "172.16.66.188 / corptest", "当前可用环境", "材料中不保存登录密码"],
    ["统一代理", f"{PROXY_DIR}/server.js", "Claude + Codex 共同入口", "监听 127.0.0.1:38443"],
    ["LaunchAgent", f"{REMOTE_HOME}/Library/LaunchAgents/com.cj.claude-local-https-proxy.plist", "登录后自动拉起代理", "stdout/stderr 写入 proxy logs"],
    ["Claude Desktop 3P", f"{REMOTE_HOME}/Library/Application Support/Claude-3p/configLibrary/*.json", "Desktop Gateway 配置", "base URL 指向 38443"],
    ["Claude CLI settings", f"{REMOTE_HOME}/.claude/settings.json", "Claude Code host/CLI 配置", "保留 Claude 槽位模型名"],
    ["Codex config", f"{REMOTE_HOME}/.codex/config.toml", "Codex provider + profiles", "provider 为 ark-coding"],
    ["旧 Codex 代理", f"{REMOTE_HOME}/.codex/ark-coding-proxy/server.js", "回滚参考", "38444 当前不应监听"],
]


ROUTES = [
    ["/health", "GET", "本机健康检查", "返回 Claude upstream、Codex upstream 与模型映射"],
    ["/healthz", "GET", "轻量健康检查", "返回 ok"],
    ["*/responses", "POST", "Codex Responses API", "转换为 Chat Completions 再请求 Ark coding v3"],
    ["其他路径", "任意", "Claude Anthropic-compatible API", "透传请求并做 Claude 槽位模型映射"],
]


MODEL_TABLE = [
    ["工具", "客户端模型名", "代理/上游模型名", "用途建议"],
    ["Claude", "claude-opus-4-6", "glm-5.1", "复杂推理/高质量输出"],
    ["Claude", "claude-sonnet-4-6", "kimi-k2.6", "默认主力模型"],
    ["Claude", "claude-haiku-4-5", "doubao-seed-2.0-pro", "快速/低成本任务"],
    ["Codex profile ark-doubao", "doubao-seed-2.0-pro", "doubao-seed-2.0-pro", "默认与快速任务"],
    ["Codex profile ark-kimi", "kimi-k2.6", "kimi-k2.6", "编码与复杂修改"],
    ["Codex profile ark-glm", "glm-5.1", "glm-5.1", "高质量推理任务"],
]


VALIDATION = [
    ["端口", "lsof 显示 127.0.0.1:38443 LISTEN", "通过"],
    ["旧端口", "127.0.0.1:38444 不再 LISTEN", "通过"],
    ["健康检查", "curl -sk https://127.0.0.1:38443/health", "通过"],
    ["Claude", "/v1/messages 返回 200，日志出现槽位映射", "通过"],
    ["Codex default", "默认 doubao-seed-2.0-pro 能回复", "通过"],
    ["Codex profiles", "ark-doubao / ark-kimi / ark-glm 均能回复", "通过"],
    ["Tool call", "Codex tool call pwd 等测试可用", "通过"],
]


CODEX_CONFIG = """model_provider = "ark-coding"
model = "doubao-seed-2.0-pro"
model_reasoning_effort = "medium"
disable_response_storage = true

[model_providers.ark-coding]
name = "Volcengine Ark Coding via unified local proxy"
wire_api = "responses"
requires_openai_auth = true
base_url = "https://127.0.0.1:38443/v1"
supports_websockets = false

[profiles.ark-doubao]
model_provider = "ark-coding"
model = "doubao-seed-2.0-pro"
model_reasoning_effort = "medium"

[profiles.ark-kimi]
model_provider = "ark-coding"
model = "kimi-k2.6"
model_reasoning_effort = "high"

[profiles.ark-glm]
model_provider = "ark-coding"
model = "glm-5.1"
model_reasoning_effort = "high"
"""


OPERATIONS = [
    [
        "检查 LaunchAgent",
        "launchctl print gui/$(id -u)/com.cj.claude-local-https-proxy",
    ],
    [
        "重启统一代理",
        "launchctl kickstart -k gui/$(id -u)/com.cj.claude-local-https-proxy",
    ],
    [
        "检查端口",
        "lsof -nP -iTCP:38443 -sTCP:LISTEN\nlsof -nP -iTCP:38444 -sTCP:LISTEN",
    ],
    [
        "检查健康状态",
        "curl -sk https://127.0.0.1:38443/health\ncurl -sk https://127.0.0.1:38443/healthz",
    ],
    [
        "查看代理日志",
        f"tail -n 160 {PROXY_DIR}/logs/proxy.log\n"
        f"tail -n 80 {PROXY_DIR}/logs/proxy.err.log",
    ],
    [
        "Codex profile smoke test",
        "codex -p ark-doubao\ncodex -p ark-kimi\ncodex -p ark-glm",
    ],
]


TROUBLESHOOTING = [
    ["health 不通", "先看 LaunchAgent state，再看 38443 端口和 proxy.err.log；证书文件路径错误也会导致代理无法启动。"],
    ["Claude App 显示 gateway unhealthy", "检查 Desktop 3P config base URL 是否仍为 https://127.0.0.1:38443；再看 main.log 里的 ConfigHealth。"],
    ["Claude 调用到了错误模型", "检查请求模型名是否包含 opus/sonnet/haiku；代理只对 Claude 槽位名做映射。"],
    ["Codex 401/403", "检查 Codex provider token 或环境变量，不要把真实 key 写进 server.js 或文档。"],
    ["Codex tool call 异常", "重点看 /responses 转换：function_call、function_call_output、tools 参数是否被正确转换。"],
    ["上游超时", "确认公司网络到 ark.cn-beijing.volces.com 可达；必要时调大 UPSTREAM_TIMEOUT_MS。"],
]


ROLLBACK = [
    "停止当前统一代理：`launchctl bootout gui/$(id -u) ~/Library/LaunchAgents/com.cj.claude-local-https-proxy.plist`。",
    "恢复代理备份：`server.js.bak.codex-merge.20260514144040`。",
    "恢复 Codex 配置备份：`~/.codex/config.toml.bak.unified.20260514145435` 或 `~/.codex/config.toml.bak.real-profiles.20260514150226`。",
    "如必须回到双代理模式，再重新加载旧 `com.cj.codex-ark-coding-proxy.plist` 并确认 38444 监听。",
    "回滚后分别跑 Claude 与 Codex smoke test；不要只看端口存在。",
]


SECURITY = [
    "文档、PPT、runbook 不保存真实 API key、SSH 密码、私钥内容。",
    "私钥文件如 `certs/server.key` 只保留在目标 Mac；迁移时优先重新生成证书。",
    "代理日志公开前先检查 authorization、cookie、个人路径等敏感内容。",
    "API key 不写死在 `server.js`；优先由客户端配置或环境变量提供。",
    "远程 SSH 密码应在完成配置后轮换或替换为 SSH key。",
]


def md_table(headers: list[str], rows: list[list[str]]) -> str:
    lines = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join(["---"] * len(headers)) + " |",
    ]
    for row in rows:
        lines.append("| " + " | ".join(cell.replace("\n", "<br>") for cell in row) + " |")
    return "\n".join(lines)


def build_markdown() -> str:
    lines: list[str] = [
        f"# {TITLE}",
        "",
        f"{SUBTITLE}",
        "",
        f"更新时间：{DATE}",
        "",
        "> 本文基于原 `handoff.md`、远端 Mac 当前状态与本轮验证结果整理；不包含真实 API key、密码或私钥。",
        "",
        "## 1. Handoff 摘要",
        "",
    ]
    lines.extend(f"- {item}" for item in SUMMARY)
    lines.extend(
        [
            "",
            "## 2. 当前架构",
            "",
            "```text",
            ARCHITECTURE.rstrip(),
            "```",
            "",
            "### 2.1 组件清单",
            "",
            md_table(["组件", "位置", "作用", "备注"], INVENTORY),
            "",
            "### 2.2 代理路由",
            "",
            md_table(["路径", "方法", "服务对象", "行为"], ROUTES),
            "",
            "## 3. 模型策略",
            "",
            "Claude 侧保留槽位模型名，便于 Desktop / CLI 兼容；Codex 侧直接使用真实模型名，避免再做一层 Claude-style 映射。",
            "",
            md_table(MODEL_TABLE[0], MODEL_TABLE[1:]),
            "",
            "## 4. Codex 多模型配置",
            "",
            "Codex 通过 provider + profiles 实现多模型切换。关键配置如下，真实 API key 不写入本文。",
            "",
            "```toml",
            CODEX_CONFIG.rstrip(),
            "```",
            "",
            "常用切换方式：",
            "",
            "```bash",
            "codex -p ark-doubao",
            "codex -p ark-kimi",
            "codex -p ark-glm",
            "```",
            "",
            "## 5. 已验证结果",
            "",
            md_table(["检查项", "信号", "状态"], VALIDATION),
            "",
            "## 6. 日常操作",
            "",
        ]
    )
    for idx, (name, command) in enumerate(OPERATIONS, start=1):
        lines.extend([f"### 6.{idx} {name}", "", "```bash", command.rstrip(), "```", ""])
    lines.extend(["## 7. 故障排查", "", md_table(["问题", "处理建议"], TROUBLESHOOTING), ""])
    lines.extend(["## 8. 回滚方案", ""])
    lines.extend(f"{idx}. {item}" for idx, item in enumerate(ROLLBACK, start=1))
    lines.extend(["", "## 9. 安全与分享边界", ""])
    lines.extend(f"- {item}" for item in SECURITY)
    lines.extend(
        [
            "",
            "## 10. 后续建议",
            "",
            "- 把远端合并后的 `server.js` 同步回材料仓库，避免本地源码与目标 Mac 实际状态分叉。",
            "- 为统一代理补一个最小自动化 smoke test：health、Claude slot mapping、Codex profiles、tool call。",
            "- 把安装动作脚本化：证书生成/信任、LaunchAgent、Claude config、Codex config、回滚备份。",
        ]
    )
    return "\n".join(lines) + "\n"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_run_font(run, east_asia: str = "PingFang SC", latin: str = "Aptos") -> None:
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(9.5)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")

    for name, size, color in [
        ("Heading 1", 18, "10202A"),
        ("Heading 2", 13, "176B87"),
        ("Heading 3", 10.5, "2A9D8F"),
    ]:
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")

    code_style = styles.add_style("Runbook Code", WD_STYLE_TYPE.CHARACTER)
    code_style.font.name = "Menlo"
    code_style.font.size = Pt(7.8)
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Menlo")


def add_para(doc: Document, text: str = "", style: str | None = None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        set_run_font(run)


def add_code(doc: Document, value: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.12)
    set_cell = OxmlElement("w:shd")
    set_cell.set(qn("w:fill"), "F3F7FA")
    p._p.get_or_add_pPr().append(set_cell)
    run = p.add_run(value.rstrip())
    run.style = "Runbook Code"
    run.font.color.rgb = RGBColor.from_string("10202A")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
        set_cell_shading(hdr[idx], "EAF5F7")
        for paragraph in hdr[idx].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor.from_string("10202A")
                set_run_font(run)
    for row in rows:
        cells = table.add_row().cells
        for idx, cell_text in enumerate(row):
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            cells[idx].text = cell_text
            for paragraph in cells[idx].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(7.8)
                    set_run_font(run)
    doc.add_paragraph()


def add_title_page(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run(TITLE)
    set_run_font(r, latin="Aptos Display")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor.from_string("10202A")

    sub = doc.add_paragraph()
    r = sub.add_run(SUBTITLE)
    set_run_font(r)
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string("60717E")

    meta = doc.add_paragraph()
    r = meta.add_run(f"更新时间：{DATE} | 目标 Mac：172.16.66.188 | 文档不含真实密钥")
    set_run_font(r)
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string("60717E")


def build_docx() -> None:
    doc = Document()
    configure_document(doc)
    add_title_page(doc)

    doc.add_heading("1. Handoff 摘要", level=1)
    add_bullets(doc, SUMMARY)

    doc.add_heading("2. 当前架构", level=1)
    add_code(doc, ARCHITECTURE)
    doc.add_heading("2.1 组件清单", level=2)
    add_table(doc, ["组件", "位置", "作用", "备注"], INVENTORY)
    doc.add_heading("2.2 代理路由", level=2)
    add_table(doc, ["路径", "方法", "服务对象", "行为"], ROUTES)

    doc.add_heading("3. 模型策略", level=1)
    add_para(doc, "Claude 侧保留槽位模型名，Codex 侧直接使用真实模型名。这样 Claude Desktop 的兼容层仍稳定，Codex 的 profile 也更直观。")
    add_table(doc, MODEL_TABLE[0], MODEL_TABLE[1:])

    doc.add_heading("4. Codex 多模型配置", level=1)
    add_para(doc, "关键配置如下，真实 API key 不写入文档。")
    add_code(doc, CODEX_CONFIG)
    add_para(doc, "常用命令：")
    add_code(doc, "codex -p ark-doubao\ncodex -p ark-kimi\ncodex -p ark-glm")

    doc.add_heading("5. 已验证结果", level=1)
    add_table(doc, ["检查项", "信号", "状态"], VALIDATION)

    doc.add_heading("6. 日常操作", level=1)
    for name, command in OPERATIONS:
        doc.add_heading(name, level=2)
        add_code(doc, command)

    doc.add_heading("7. 故障排查", level=1)
    add_table(doc, ["问题", "处理建议"], TROUBLESHOOTING)

    doc.add_heading("8. 回滚方案", level=1)
    add_numbered(doc, ROLLBACK)

    doc.add_heading("9. 安全与分享边界", level=1)
    add_bullets(doc, SECURITY)

    doc.add_heading("10. 后续建议", level=1)
    add_bullets(
        doc,
        [
            "把远端合并后的 server.js 同步回材料仓库，避免本地源码与目标 Mac 实际状态分叉。",
            "为统一代理补一个最小自动化 smoke test：health、Claude slot mapping、Codex profiles、tool call。",
            "把安装动作脚本化：证书生成/信任、LaunchAgent、Claude config、Codex config、回滚备份。",
        ],
    )

    doc.add_section(WD_SECTION.CONTINUOUS)
    doc.save(DOCX_PATH)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    MD_PATH.write_text(build_markdown(), encoding="utf-8")
    build_docx()
    print(MD_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
