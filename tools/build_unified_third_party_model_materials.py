#!/usr/bin/env python3
"""Build consolidated Word manual and AI runbook for Claude + Codex unified proxy."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs"

MANUAL_DOCX = OUT / "unified-proxy-third-party-models-technical-manual.docx"
MANUAL_MD = OUT / "unified-proxy-third-party-models-technical-manual.md"
AI_RUNBOOK_MD = OUT / "unified-proxy-third-party-models-ai-runbook.md"

TITLE = "统一代理接入第三方模型技术手册"
SUBTITLE = "Claude Code 与 Codex 通过本机 HTTPS 代理统一接入 OpenAI-compatible / Anthropic-compatible 模型服务"
DATE = "2026-05-14"

PROXY_URL = "https://127.0.0.1:38443"
CLAUDE_UPSTREAM = "https://ark.cn-beijing.volces.com/api/coding"
CODEX_UPSTREAM = "https://ark.cn-beijing.volces.com/api/coding/v3"


def table_md(headers: list[str], rows: list[list[str]]) -> str:
    lines = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join(["---"] * len(headers)) + " |",
    ]
    for row in rows:
        lines.append("| " + " | ".join(cell.replace("\n", "<br>") for cell in row) + " |")
    return "\n".join(lines)


ARCHITECTURE = """Claude Code Desktop / Claude Code CLI
  Anthropic-compatible API
  /v1/messages, /v1/messages/count_tokens
                       |
                       | HTTPS + local trusted CA
                       v
            https://127.0.0.1:38443
             claude-local-proxy/server.js
              |                       |
              | passthrough +         | Responses API bridge
              | Claude model mapping  |
              v                       v
  https://ark.cn-beijing.volces.com/api/coding
  https://ark.cn-beijing.volces.com/api/coding/v3/chat/completions
                                      ^
                                      |
                                Codex CLI / App
                                /v1/responses
"""

COMPONENTS = [
    ["本机统一代理", "<PROJECT_ROOT>/claude-local-proxy/server.js", "唯一入口；监听 127.0.0.1:38443；处理 Claude 透传与 Codex 协议转换"],
    ["本地证书", "<PROJECT_ROOT>/claude-local-proxy/certs/ca.crt, server.crt, server.key", "让 Desktop/Electron、CLI、host loop 都能信任本机 HTTPS"],
    ["LaunchAgent", "~/Library/LaunchAgents/com.cj.claude-local-https-proxy.plist", "登录后自动启动代理；提供上游地址、模型映射和证书路径"],
    ["Claude Desktop 3P config", "~/Library/Application Support/Claude-3p/configLibrary/*.json", "Desktop third-party inference 指向统一代理"],
    ["Claude CLI settings", "~/.claude/settings.json", "Claude Code CLI/host 通过 ANTHROPIC_* 环境变量指向统一代理"],
    ["Codex config", "~/.codex/config.toml", "Codex provider 使用 responses wire API，profiles 直接使用真实模型名"],
    ["验证日志", "<PROJECT_ROOT>/claude-local-proxy/logs/*.log, ~/Library/Logs/Claude-3p/main.log", "排查入口健康、模型映射、协议转换与上游错误"],
]

MODEL_ROWS = [
    ["Claude Code", "claude-opus-4-6", "glm-5.1", "通过代理映射", "复杂推理/高质量输出"],
    ["Claude Code", "claude-sonnet-4-6", "kimi-k2.6", "通过代理映射", "默认主力模型"],
    ["Claude Code", "claude-haiku-4-5", "doubao-seed-2.0-pro", "通过代理映射", "快速/低成本任务"],
    ["Codex ark-doubao", "doubao-seed-2.0-pro", "doubao-seed-2.0-pro", "真实模型名", "默认/快速任务"],
    ["Codex ark-kimi", "kimi-k2.6", "kimi-k2.6", "真实模型名", "复杂编码任务"],
    ["Codex ark-glm", "glm-5.1", "glm-5.1", "真实模型名", "高质量推理任务"],
]

ROUTE_ROWS = [
    ["/health", "GET", "运维健康检查", "返回 ok、Claude upstream、Codex upstream、模型映射"],
    ["/healthz", "GET", "轻量健康检查", "返回 ok"],
    ["*/responses", "POST", "Codex CLI/App", "解析 Responses API，转换为 Chat Completions，再还原 Responses 输出"],
    ["其他路径", "任意", "Claude Code Desktop/CLI", "转发到 Anthropic-compatible upstream，并按 Claude 槽位映射模型"],
]

VALIDATION_ROWS = [
    ["代理进程", "launchctl print gui/$(id -u)/com.cj.claude-local-https-proxy", "state = running"],
    ["端口", "lsof -nP -iTCP:38443 -sTCP:LISTEN", "只有统一代理监听 38443；旧 38444 不监听"],
    ["健康检查", "curl -sk https://127.0.0.1:38443/health", "返回 ok、upstream、codexUpstream 与模型字段"],
    ["Claude Desktop", "~/Library/Logs/Claude-3p/main.log", "ConfigHealth recomputed { state: 'healthy', provider: 'gateway' }"],
    ["Claude 请求", "代理日志", "POST /v1/messages -> 200；出现 claude-* -> 真实模型映射"],
    ["Codex 默认", "codex", "默认 doubao-seed-2.0-pro 可回复"],
    ["Codex profiles", "codex -p ark-doubao / ark-kimi / ark-glm", "三个 profile 均可回复"],
    ["Tool call", "Codex 中触发 shell/tool call", "function_call 与 function_call_output 闭环正常"],
]

FAILURE_ROWS = [
    ["Desktop gateway unhealthy", "Desktop 3P config base URL 错误、证书未信任或代理未运行", "先查 /health、Keychain、main.log，再查代理日志"],
    ["Claude 调错模型", "保留了 ANTHROPIC_MODEL 或 modelOverrides，或请求模型名不是 Claude 槽位", "删除强制模型字段，只保留槽位默认值；映射放在代理层"],
    ["Codex 401/403", "API key 未进入 Codex provider 或 Authorization 未透传", "检查 ~/.codex/config.toml 与环境变量；不要把 key 写入 server.js"],
    ["Codex tool call 失败", "Responses API 与 Chat Completions 转换不完整", "检查 function_call、function_call_output、tools、tool_choice 的转换日志"],
    ["Cowork server is busy", "UI 泛化错误，真实原因常是 SSL certificate verification failed", "安装 claude-ca-launcher，给 host loop 注入 CA 环境"],
    ["System keychain 写入失败", "SSH 无交互授权被 macOS 拦截", "在目标 Mac 本机交互式 sudo，或用 MDM/配置描述文件下发 CA"],
]

ENV_SAMPLE = """LISTEN_HOST=127.0.0.1
LISTEN_PORT=38443
UPSTREAM_BASE_URL=https://ark.cn-beijing.volces.com/api/coding
CODEX_UPSTREAM_BASE_URL=https://ark.cn-beijing.volces.com/api/coding/v3
BIG_MODEL=glm-5.1
MIDDLE_MODEL=kimi-k2.6
SMALL_MODEL=doubao-seed-2.0-pro
TLS_CERT_FILE=<PROJECT_ROOT>/claude-local-proxy/certs/server.crt
TLS_KEY_FILE=<PROJECT_ROOT>/claude-local-proxy/certs/server.key
UPSTREAM_TIMEOUT_MS=300000
KEEP_ALIVE_TIMEOUT_MS=300000
HEADERS_TIMEOUT_MS=310000"""

CLAUDE_DESKTOP_JSON = """{
  "provider": "gateway",
  "gatewayBaseUrl": "https://127.0.0.1:38443",
  "gatewayApiKey": "<ARK_API_KEY>",
  "gatewayAuthScheme": "bearer",
  "inferenceModels": [
    { "id": "claude-sonnet-4-6", "name": "Sonnet 4.6" },
    { "id": "claude-opus-4-6", "name": "Opus 4.6" },
    { "id": "claude-haiku-4-5", "name": "Haiku 4.5" }
  ],
  "hideAnthropicSignIn": true
}"""

CLAUDE_SETTINGS_JSON = """{
  "env": {
    "ANTHROPIC_BASE_URL": "https://127.0.0.1:38443",
    "ANTHROPIC_AUTH_TOKEN": "<ARK_API_KEY>",
    "ANTHROPIC_DEFAULT_OPUS_MODEL": "claude-opus-4-6",
    "ANTHROPIC_DEFAULT_SONNET_MODEL": "claude-sonnet-4-6",
    "ANTHROPIC_DEFAULT_HAIKU_MODEL": "claude-haiku-4-5",
    "NODE_USE_SYSTEM_CA": "1",
    "NODE_EXTRA_CA_CERTS": "<PROJECT_ROOT>/claude-local-proxy/certs/ca.crt",
    "SSL_CERT_FILE": "<PROJECT_ROOT>/claude-local-proxy/certs/ca.crt"
  }
}"""

CODEX_TOML = """model_provider = "ark-coding"
model = "doubao-seed-2.0-pro"
model_reasoning_effort = "medium"
disable_response_storage = true

[model_providers.ark-coding]
name = "Volcengine Ark Coding via unified local proxy"
wire_api = "responses"
requires_openai_auth = true
base_url = "https://127.0.0.1:38443/v1"
supports_websockets = false

[profiles.ark-doubao]
model_provider = "ark-coding"
model = "doubao-seed-2.0-pro"
model_reasoning_effort = "medium"

[profiles.ark-kimi]
model_provider = "ark-coding"
model = "kimi-k2.6"
model_reasoning_effort = "high"

[profiles.ark-glm]
model_provider = "ark-coding"
model = "glm-5.1"
model_reasoning_effort = "high"
"""

CODEX_APP_OVERRIDE_COMMANDS = """# 以 Kimi profile 对应模型打开某个 workspace
codex app /path/to/project \\
  -c 'model_provider="ark-coding"' \\
  -c 'model="kimi-k2.6"' \\
  -c 'model_reasoning_effort="high"'

# 以 Doubao 打开某个 workspace
codex app /path/to/project \\
  -c 'model_provider="ark-coding"' \\
  -c 'model="doubao-seed-2.0-pro"' \\
  -c 'model_reasoning_effort="medium"'

# 以 GLM 打开某个 workspace
codex app /path/to/project \\
  -c 'model_provider="ark-coding"' \\
  -c 'model="glm-5.1"' \\
  -c 'model_reasoning_effort="high"'"""

CODEX_APP_SWITCH_ROWS = [
    ["Codex CLI", "`codex -p ark-kimi`", "读取 `[profiles.*]`，适合命令行临时切换。"],
    ["Codex App 默认", "改 `~/.codex/config.toml` 顶层 `model` 后重启 App", "最稳方式；新会话按默认模型启动。"],
    ["Codex App 临时", "`codex app /path -c 'model=\"kimi-k2.6\"' ...`", "启动 App 时覆盖配置；适合一次性打开特定 workspace。"],
]


MANUAL_SECTIONS = [
    ("1. 目的与适用范围", [
        "本文定义一套标准方案：通过一个本机 HTTPS 统一代理，让 Claude Code Desktop/CLI 与 Codex CLI/App 同时接入第三方模型服务。",
        "适用对象包括个人 Mac、公司内网 Mac、远程 SSH 管理的 macOS 机器，以及需要把 OpenAI-compatible 或 Anthropic-compatible 服务接入本地 AI 编程工具的场景。",
        "本文使用火山 Ark coding endpoint 作为已验证示例；其他兼容服务可以沿用相同结构，只需要替换 upstream、模型名与认证方式。",
    ]),
    ("2. 统一代理优势", [
        "单一入口：Claude Code 与 Codex 都指向本机 HTTPS endpoint，端口、证书、日志和启动方式统一管理。",
        "清晰模型策略：Claude Code 侧保留 Claude 槽位名，Codex 侧直接使用真实模型名，避免两套工具互相影响。",
        "协议差异内聚：Claude 的 Anthropic-compatible 请求与 Codex 的 Responses API 请求都在代理层适配，上游服务只承担模型推理。",
        "验证口径一致：health、代理日志、Desktop main.log、Codex profile smoke test 与 tool call 构成统一验收矩阵。",
        "安全边界清楚：API key 只进入本机配置或环境变量，不写入代理源码、文档、PPT、runbook 或 Git。",
    ]),
]


def manual_markdown() -> str:
    lines = [
        f"# {TITLE}",
        "",
        SUBTITLE,
        "",
        f"版本日期：{DATE}",
        "",
        "> 安全说明：本文只使用 `<ARK_API_KEY>`、`<PROJECT_ROOT>` 等占位符，不包含真实 API key、SSH 密码或私钥。",
        "",
    ]
    for heading, paragraphs in MANUAL_SECTIONS:
        lines.extend([f"## {heading}", ""])
        lines.extend([p + "\n" for p in paragraphs])

    lines.extend([
        "## 3. 目标架构",
        "",
        "```text",
        ARCHITECTURE.rstrip(),
        "```",
        "",
        "### 3.1 组件清单",
        "",
        table_md(["组件", "标准位置", "职责"], COMPONENTS),
        "",
        "### 3.2 代理路由",
        "",
        table_md(["路径", "方法", "调用方", "代理行为"], ROUTE_ROWS),
        "",
        "## 4. 模型策略",
        "",
        "Claude Code 使用槽位模型名，Codex 使用真实模型名。这是本方案最重要的配置分界。",
        "",
        table_md(["工具/配置", "客户端模型名", "实际上游模型", "策略", "建议用途"], MODEL_ROWS),
        "",
        "## 5. 统一代理配置",
        "",
        "LaunchAgent 或 shell 环境应提供以下变量。`CODEX_UPSTREAM_BASE_URL` 指向 OpenAI-compatible coding v3，`UPSTREAM_BASE_URL` 指向 Claude/Anthropic-compatible endpoint。",
        "",
        "```bash",
        ENV_SAMPLE,
        "```",
        "",
        "代理必须至少实现两条逻辑：",
        "",
        "1. Claude 分支：透传 `/v1/messages`、`/v1/messages/count_tokens` 等 Anthropic-compatible 请求，并把 Claude 槽位名改写为真实模型。",
        "2. Codex 分支：接收 `/v1/responses`，转换成 `/chat/completions`，再把上游结果还原为 Responses API 输出。",
        "",
        "## 6. 配置 Claude Code",
        "",
        "### 6.1 Claude Desktop 3P Gateway",
        "",
        "```json",
        CLAUDE_DESKTOP_JSON,
        "```",
        "",
        "### 6.2 Claude Code CLI / host settings",
        "",
        "```json",
        CLAUDE_SETTINGS_JSON,
        "```",
        "",
        "配置要求：删除 `ANTHROPIC_MODEL`，删除旧 `modelOverrides`。默认模型只写 Claude 槽位名，真实模型映射放在代理层。",
        "",
        "### 6.3 Desktop host binary 与 Cowork",
        "",
        "Desktop Code host 的版本目录以 `~/Library/Logs/Claude-3p/main.log` 中 `[CCD] Initialized with version ...` 为准。若 Cowork 报 SSL 证书失败，需要通过 launcher 注入 `NODE_USE_SYSTEM_CA`、`NODE_EXTRA_CA_CERTS` 与 `SSL_CERT_FILE`。",
        "",
        "## 7. 配置 Codex",
        "",
        "Codex 通过 provider + profiles 使用统一代理。`base_url` 指向 `https://127.0.0.1:38443/v1`，`wire_api` 使用 `responses`。",
        "",
        "```toml",
        CODEX_TOML.rstrip(),
        "```",
        "",
        "常用命令：",
        "",
        "```bash",
        "codex -p ark-doubao",
        "codex -p ark-kimi",
        "codex -p ark-glm",
        "```",
        "",
        "### 7.1 Codex CLI 与 Codex App 切换模型",
        "",
        "Codex CLI 支持 `-p/--profile`，但 `codex app` 当前更适合通过默认配置或启动参数覆盖模型。不要把 CLI profile 切换方式直接套到 App 上。",
        "",
        table_md(["场景", "操作方法", "说明"], CODEX_APP_SWITCH_ROWS),
        "",
        "Codex App 默认切换：修改 `~/.codex/config.toml` 顶层 `model_provider`、`model`、`model_reasoning_effort`，然后退出并重新打开 Codex App；新会话会使用新的默认模型。",
        "",
        "Codex App 临时切换：",
        "",
        "```bash",
        CODEX_APP_OVERRIDE_COMMANDS,
        "```",
        "",
        "注意：已打开的当前会话不保证热切换模型；需要稳定验证时，新开 workspace 或重启 App 后再检查代理日志里的 `codex responses model ... -> ...`。",
        "",
        "## 8. 验证与验收",
        "",
        table_md(["检查项", "命令/位置", "通过标准"], VALIDATION_ROWS),
        "",
        "## 9. 运维与故障处理",
        "",
        table_md(["故障", "常见原因", "处理方式"], FAILURE_ROWS),
        "",
        "## 10. 回滚方案",
        "",
        "1. 停止统一代理 LaunchAgent。",
        "2. 恢复 `server.js` 备份。",
        "3. 恢复 `~/.codex/config.toml` 与 `~/.claude/settings.json` 备份。",
        "4. 如需恢复双代理模式，再加载旧 Codex LaunchAgent 并确认旧端口监听。",
        "5. 回滚后分别验证 Claude `/v1/messages`、Codex `/v1/responses` 与 tool call。",
        "",
        "## 11. 安全与交付边界",
        "",
        "- 不在材料中保存真实 API key、SSH 密码、cookie、私钥或完整 authorization header。",
        "- `certs/*.key` 不提交公共仓库；迁移机器时优先重新生成证书。",
        "- 代理日志对外分享前必须扫描 `authorization`、`x-api-key`、`api-key`、`cookie`。",
        "- 远程 SSH 密码在完成配置后建议替换为 SSH key 或轮换。",
        "- 公司批量部署时，CA 信任优先通过 MDM/配置描述文件下发。",
    ])
    return "\n".join(lines) + "\n"


AI_RUNBOOK = f"""# AI 工具执行 Runbook：通过统一代理配置 Claude Code 与 Codex 第三方模型

版本日期：{DATE}

## 0. 任务目标

在一台 macOS 上配置一个本机 HTTPS 统一代理，让：

- Claude Code Desktop / Claude Code CLI 通过 `{PROXY_URL}` 使用第三方 Anthropic-compatible 模型服务。
- Codex CLI / App 通过同一个 `{PROXY_URL}` 使用第三方 OpenAI-compatible Responses API 入口。
- Claude 侧保留 Claude 槽位模型名；Codex 侧直接使用真实模型名。

## 1. 输入参数

执行前向用户确认或从环境中读取以下参数：

| 参数 | 示例 | 说明 |
| --- | --- | --- |
| `<HOST>` | `172.16.x.x` | 目标 Mac，可为本机或 SSH 主机 |
| `<USER>` | `corptest` | 目标 Mac 用户 |
| `<PROJECT_ROOT>` | `~/Documents/Codex/claude-code-app-api` | 代理项目目录 |
| `<NODE_BIN>` | `/usr/local/bin/node` | LaunchAgent 使用的 Node 路径 |
| `<ARK_API_KEY>` | 不写入文档 | 真实 key 只能进入本机配置或环境变量 |
| `<CLAUDE_UPSTREAM>` | `{CLAUDE_UPSTREAM}` | Anthropic-compatible endpoint |
| `<CODEX_UPSTREAM>` | `{CODEX_UPSTREAM}` | OpenAI-compatible endpoint |

## 2. 不可违反的约束

- 不要把真实 API key、密码、cookie、私钥写入 Markdown、Word、PPT、Git commit、日志摘要或最终回复。
- 修改配置前先备份：`server.js`、`~/.claude/settings.json`、`~/.codex/config.toml`、LaunchAgent plist。
- 不要清空用户已有配置；只修改本任务相关字段。
- 证书私钥只保存在目标机器本地；迁移机器时优先重新生成。
- macOS System keychain 信任可能需要本机交互式 sudo，SSH 下失败时不要反复重试。

## 3. 标准模型策略

| 工具 | 客户端模型名 | 上游模型名 | 处理方式 |
| --- | --- | --- | --- |
| Claude Code | `claude-opus-4-6` | `glm-5.1` | 代理映射 |
| Claude Code | `claude-sonnet-4-6` | `kimi-k2.6` | 代理映射 |
| Claude Code | `claude-haiku-4-5` | `doubao-seed-2.0-pro` | 代理映射 |
| Codex | `doubao-seed-2.0-pro` | `doubao-seed-2.0-pro` | profile 直连 |
| Codex | `kimi-k2.6` | `kimi-k2.6` | profile 直连 |
| Codex | `glm-5.1` | `glm-5.1` | profile 直连 |

## 4. 执行步骤

### Step 1：采集现状

```bash
uname -m
command -v node || true
command -v claude || true
command -v codex || true
lsof -nP -iTCP:38443 -sTCP:LISTEN || true
lsof -nP -iTCP:38444 -sTCP:LISTEN || true
```

记录：

- 当前是否已有 38443 代理。
- 是否存在旧 38444 Codex 代理。
- Claude Desktop 是否有 `~/Library/Application Support/Claude-3p`。
- Codex 是否有 `~/.codex/config.toml`。

### Step 2：安装或更新统一代理

代理必须支持：

- `GET /health`
- `GET /healthz`
- Claude passthrough：`/v1/messages`、`/v1/messages/count_tokens`
- Codex bridge：`POST */responses` -> `{CODEX_UPSTREAM}/chat/completions`
- Claude slot mapping：opus/sonnet/haiku -> glm/kimi/doubao

启动环境变量：

```bash
{ENV_SAMPLE}
```

### Step 3：生成并信任本机证书

证书要求：

- server certificate SAN 包含 `127.0.0.1`、`localhost`，建议也包含 `::1`。
- Desktop/Electron 需要 Keychain 信任，不要只依赖 curl `--cacert`。
- 如果 System keychain 无法通过 SSH 写入，让用户在目标 Mac 本机执行。

验证：

```bash
curl --cacert <PROJECT_ROOT>/claude-local-proxy/certs/ca.crt https://127.0.0.1:38443/health
curl -sk https://127.0.0.1:38443/health
```

### Step 4：配置 LaunchAgent

目标文件：

```text
~/Library/LaunchAgents/com.cj.claude-local-https-proxy.plist
```

加载与验证：

```bash
launchctl unload ~/Library/LaunchAgents/com.cj.claude-local-https-proxy.plist 2>/dev/null || true
launchctl load ~/Library/LaunchAgents/com.cj.claude-local-https-proxy.plist
launchctl kickstart -k gui/$(id -u)/com.cj.claude-local-https-proxy
launchctl print gui/$(id -u)/com.cj.claude-local-https-proxy
```

### Step 5：配置 Claude Code

Desktop 3P Gateway：

```json
{CLAUDE_DESKTOP_JSON}
```

CLI/host settings：

```json
{CLAUDE_SETTINGS_JSON}
```

检查项：

- `ANTHROPIC_BASE_URL` 必须是 `{PROXY_URL}`。
- 不应保留 `ANTHROPIC_MODEL`。
- 不应保留旧 `modelOverrides`。
- `ANTHROPIC_DEFAULT_*_MODEL` 写 Claude 槽位名，不写真实模型名。

### Step 6：配置 Codex

目标文件：

```text
~/.codex/config.toml
```

标准配置：

```toml
{CODEX_TOML.rstrip()}
```

检查项：

- provider 名为 `ark-coding`。
- `wire_api = "responses"`。
- `base_url = "https://127.0.0.1:38443/v1"`。
- profiles 使用真实模型名。

### Step 6.1：Codex CLI 与 Codex App 切换模型

Codex CLI 用 profile 切换：

```bash
codex -p ark-doubao
codex -p ark-kimi
codex -p ark-glm
```

Codex App 推荐两种方式：

1. 修改 `~/.codex/config.toml` 顶层默认模型，然后退出并重新打开 Codex App。适合长期默认切换。
2. 启动 App 时用 `-c key=value` 临时覆盖。适合一次性打开指定 workspace。

```bash
{CODEX_APP_OVERRIDE_COMMANDS}
```

注意：当前已打开会话不保证热切换模型。验证时新开会话或重启 App，再看代理日志中的 `codex responses model <客户端模型> -> <上游模型>`。

### Step 7：验证

```bash
lsof -nP -iTCP:38443 -sTCP:LISTEN
curl -sk https://127.0.0.1:38443/health
tail -n 160 <PROJECT_ROOT>/claude-local-proxy/logs/proxy.log
codex -p ark-doubao
codex -p ark-kimi
codex -p ark-glm
```

Claude Desktop 额外看：

```bash
tail -n 200 "$HOME/Library/Logs/Claude-3p/main.log"
```

成功信号：

- `ConfigHealth recomputed {{ state: 'healthy', provider: 'gateway' }}`
- `POST /v1/messages -> 200`
- `codex responses model doubao-seed-2.0-pro -> doubao-seed-2.0-pro`
- `codex responses model kimi-k2.6 -> kimi-k2.6`
- `codex responses model glm-5.1 -> glm-5.1`
- Codex tool call 能完成 `function_call` / `function_call_output` 往返。

## 5. 故障决策树

1. `/health` 不通：先查 LaunchAgent state、38443 端口、`proxy.err.log`。
2. `/health` 通但 Claude unhealthy：查 Keychain 信任、Desktop 3P config、`main.log`。
3. Claude 模型不对：查 `ANTHROPIC_MODEL`、`modelOverrides`、代理映射日志。
4. Codex 不通：查 `~/.codex/config.toml`、Authorization、`/v1/responses` 分支日志。
5. Codex tool call 不通：查 Responses bridge 的 tool 转换逻辑。
6. Cowork 显示 server busy：查 transcript 真实错误；若是 SSL，安装 `claude-ca-launcher`。

## 6. 回滚步骤

```bash
launchctl bootout gui/$(id -u) ~/Library/LaunchAgents/com.cj.claude-local-https-proxy.plist 2>/dev/null || true
```

然后按备份恢复：

- `<PROJECT_ROOT>/claude-local-proxy/server.js.bak.*`
- `~/.claude/settings.json.bak.*`
- `~/.codex/config.toml.bak.*`
- 如需双代理模式，恢复旧 Codex LaunchAgent，并确认 38444 重新监听。

回滚后仍要分别验证 Claude 和 Codex，不要只看进程存在。

## 7. 交付输出格式

执行完成后，给用户输出：

- 修改了哪些文件。
- 当前监听端口。
- Claude 和 Codex 各自成功信号。
- profiles 是否通过。
- 是否存在未完成的证书信任、Cowork 或网络下载问题。
- 所有敏感值均以 `<REDACTED>` 或占位符呈现。
"""


def set_run_font(run, east_asia: str = "PingFang SC", latin: str = "Calibri") -> None:
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color.replace("#", ""))


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.replace("#", ""))
    tc_pr.append(shd)


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title_style = styles["Title"]
    title_style.font.name = "Calibri Light"
    title_style.font.size = Pt(24)
    title_style.font.color.rgb = rgb("0B2545")
    title_style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    code_style = styles.add_style("Manual Code", WD_STYLE_TYPE.CHARACTER)
    code_style.font.name = "Menlo"
    code_style.font.size = Pt(7.4)
    code_style.font.color.rgb = rgb("10202A")
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Menlo")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("统一代理第三方模型技术手册")
    set_run_font(r)
    r.font.size = Pt(8)
    r.font.color.rgb = rgb("60717E")


def paragraph(doc: Document, text: str = "", style: str | None = None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    set_run_font(r)
    return p


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r)


def numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r)


def code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.08)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F3F7FA")
    p._p.get_or_add_pPr().append(shd)
    r = p.add_run(text.rstrip())
    r.style = "Manual Code"


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        shade_cell(hdr[i], "F2F4F7")
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            hdr[i].width = Inches(widths[i])
        for p in hdr[i].paragraphs:
            p.paragraph_format.space_after = Pt(0)
            for r in p.runs:
                set_run_font(r)
                r.bold = True
                r.font.size = Pt(8)
                r.font.color.rgb = rgb("0B2545")

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths:
                cells[i].width = Inches(widths[i])
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.08
                for r in p.runs:
                    set_run_font(r)
                    r.font.size = Pt(7.4)
                    r.font.color.rgb = rgb("10202A")
    doc.add_paragraph()


def add_callout(doc: Document, label: str, body: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.08)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "EAF5F7")
    p._p.get_or_add_pPr().append(shd)
    r1 = p.add_run(label + "：")
    set_run_font(r1)
    r1.bold = True
    r1.font.color.rgb = rgb("176B87")
    r2 = p.add_run(body)
    set_run_font(r2)


def build_docx() -> None:
    doc = Document()
    configure_doc(doc)

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = p.add_run(TITLE)
    set_run_font(title_run, latin="Calibri Light")

    sub = paragraph(doc, SUBTITLE)
    sub.runs[0].font.size = Pt(11)
    sub.runs[0].font.color.rgb = rgb("60717E")

    meta = paragraph(doc, f"版本日期：{DATE} | 密级建议：内部技术资料 | 敏感值均使用占位符")
    meta.runs[0].font.size = Pt(8.5)
    meta.runs[0].font.color.rgb = rgb("60717E")

    add_callout(
        doc,
        "核心结论",
        "Claude Code 与 Codex 可以通过同一个本机 HTTPS 代理接入第三方模型。Claude 使用槽位映射，Codex 使用真实模型名 profiles。",
    )

    for heading, paragraphs in MANUAL_SECTIONS:
        doc.add_heading(heading, level=1)
        for item in paragraphs:
            bullet(doc, item)

    doc.add_heading("3. 目标架构", level=1)
    code_block(doc, ARCHITECTURE)
    doc.add_heading("3.1 组件清单", level=2)
    add_table(doc, ["组件", "标准位置", "职责"], COMPONENTS, [1.35, 2.65, 2.3])
    doc.add_heading("3.2 代理路由", level=2)
    add_table(doc, ["路径", "方法", "调用方", "代理行为"], ROUTE_ROWS, [1.35, 0.65, 1.35, 2.95])

    doc.add_heading("4. 模型策略", level=1)
    paragraph(doc, "Claude Code 使用槽位模型名，Codex 使用真实模型名。这是本方案最重要的配置分界。")
    add_table(doc, ["工具/配置", "客户端模型名", "实际上游模型", "策略", "建议用途"], MODEL_ROWS, [1.2, 1.55, 1.35, 1.05, 1.15])

    doc.add_heading("5. 统一代理配置", level=1)
    paragraph(doc, "LaunchAgent 或 shell 环境应提供以下变量。")
    code_block(doc, ENV_SAMPLE)
    doc.add_heading("5.1 代理必须实现的逻辑", level=2)
    numbered(doc, "Claude 分支：透传 /v1/messages、/v1/messages/count_tokens 等 Anthropic-compatible 请求，并把 Claude 槽位名改写为真实模型。")
    numbered(doc, "Codex 分支：接收 /v1/responses，转换成 /chat/completions，再把上游结果还原为 Responses API 输出。")

    doc.add_heading("6. 配置 Claude Code", level=1)
    doc.add_heading("6.1 Claude Desktop 3P Gateway", level=2)
    code_block(doc, CLAUDE_DESKTOP_JSON)
    doc.add_heading("6.2 Claude Code CLI / host settings", level=2)
    code_block(doc, CLAUDE_SETTINGS_JSON)
    add_callout(doc, "配置要求", "删除 ANTHROPIC_MODEL 和旧 modelOverrides；默认模型只写 Claude 槽位名。")
    doc.add_heading("6.3 Desktop host binary 与 Cowork", level=2)
    paragraph(doc, "Desktop Code host 的版本目录以 Claude Desktop 日志中的 [CCD] Initialized with version ... 为准。Cowork 若报 SSL 证书失败，需要通过 launcher 注入 NODE_USE_SYSTEM_CA、NODE_EXTRA_CA_CERTS 与 SSL_CERT_FILE。")

    doc.add_heading("7. 配置 Codex", level=1)
    paragraph(doc, "Codex 通过 provider + profiles 使用统一代理。base_url 指向 https://127.0.0.1:38443/v1，wire_api 使用 responses。")
    code_block(doc, CODEX_TOML)
    paragraph(doc, "常用 profile 切换命令：")
    code_block(doc, "codex -p ark-doubao\ncodex -p ark-kimi\ncodex -p ark-glm")
    doc.add_heading("7.1 Codex CLI 与 Codex App 切换模型", level=2)
    paragraph(doc, "Codex CLI 支持 -p/--profile；Codex App 更适合通过默认配置或启动参数覆盖模型。不要把 CLI profile 切换方式直接套到 App 上。")
    add_table(doc, ["场景", "操作方法", "说明"], CODEX_APP_SWITCH_ROWS, [1.25, 2.25, 2.8])
    paragraph(doc, "Codex App 默认切换：修改 ~/.codex/config.toml 顶层 model_provider、model、model_reasoning_effort，然后退出并重新打开 Codex App；新会话会使用新的默认模型。")
    paragraph(doc, "Codex App 临时切换：")
    code_block(doc, CODEX_APP_OVERRIDE_COMMANDS)
    add_callout(doc, "注意", "当前已打开会话不保证热切换模型。验证时新开会话或重启 App，再看代理日志中的 codex responses model <客户端模型> -> <上游模型>。")

    doc.add_heading("8. 验证与验收", level=1)
    add_table(doc, ["检查项", "命令/位置", "通过标准"], VALIDATION_ROWS, [1.05, 2.55, 2.7])

    doc.add_heading("9. 运维与故障处理", level=1)
    add_table(doc, ["故障", "常见原因", "处理方式"], FAILURE_ROWS, [1.65, 2.15, 2.5])

    doc.add_heading("10. 回滚方案", level=1)
    for item in [
        "停止统一代理 LaunchAgent。",
        "恢复 server.js 备份。",
        "恢复 ~/.codex/config.toml 与 ~/.claude/settings.json 备份。",
        "如需恢复双代理模式，再加载旧 Codex LaunchAgent 并确认旧端口监听。",
        "回滚后分别验证 Claude /v1/messages、Codex /v1/responses 与 tool call。",
    ]:
        numbered(doc, item)

    doc.add_heading("11. 安全与交付边界", level=1)
    for item in [
        "不在材料中保存真实 API key、SSH 密码、cookie、私钥或完整 authorization header。",
        "certs/*.key 不提交公共仓库；迁移机器时优先重新生成证书。",
        "代理日志对外分享前必须扫描 authorization、x-api-key、api-key、cookie。",
        "远程 SSH 密码在完成配置后建议替换为 SSH key 或轮换。",
        "公司批量部署时，CA 信任优先通过 MDM/配置描述文件下发。",
    ]:
        bullet(doc, item)

    doc.add_section(WD_SECTION.CONTINUOUS)
    doc.save(MANUAL_DOCX)


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    MANUAL_MD.write_text(manual_markdown(), encoding="utf-8")
    AI_RUNBOOK_MD.write_text(AI_RUNBOOK, encoding="utf-8")
    build_docx()
    print(MANUAL_MD)
    print(MANUAL_DOCX)
    print(AI_RUNBOOK_MD)


if __name__ == "__main__":
    main()
